import json
import logging
from pathlib import Path
from typing import Dict, List, Optional
import pandas as pd
from datetime import datetime
import re
from ..core.reasoningLLM import ReasoningLLMEngine
from io import StringIO
import traceback

class DeliveryGenerator:
    """交付文件生成器"""
    
    def __init__(self, work_dir: str = "work_dir", reasoning_engine: Optional[ReasoningLLMEngine] = None, logger: Optional[logging.Logger] = None):
        """初始化交付文件生成器
        
        Args:
            work_dir: 工作目录
            reasoning_engine: 推理引擎实例，用于生成内容
            logger: 可选，日志记录器实例
        """
        self.work_dir = Path(work_dir)
        self.work_dir.mkdir(parents=True, exist_ok=True)
        self.logger = logger or logging.getLogger(__name__)
        self.reasoning_engine = reasoning_engine
        
        if not self.reasoning_engine:
            self.logger.warning("未提供推理引擎实例，部分功能可能受限")
        
    
    def _load_context(self, plan_dir: Path) -> Dict:
        """加载上下文信息
        
        Args:
            plan_dir: 计划目录路径
            
        Returns:
            Dict: 包含delivery_plan和search_results的上下文字典
        """
        context = {}
        try:
            # 从work_dir下的generations/inputs目录加载文件
            input_dir = self.work_dir / "generations" / "inputs"
            
            # 加载交付计划
            plan_path = input_dir / "delivery_plan.json"
            if not plan_path.exists():
                raise FileNotFoundError(f"交付计划文件不存在: {plan_path}")
            
            with plan_path.open('r', encoding='utf-8') as f:
                context['delivery_plan'] = json.load(f)
            
            # 加载检索结果
            search_results_path = input_dir / "search_results.json"
            if not search_results_path.exists():
                raise FileNotFoundError(f"检索结果文件不存在: {search_results_path}")
            
            with search_results_path.open('r', encoding='utf-8') as f:
                context['search_results'] = json.load(f)

            return context
        
        except Exception as e:
            self.logger.error(f"加载上下文信息失败: {str(e)}")
            raise
    
    def _read_file_content(self, file_path: str, encoding: str = 'utf-8') -> Optional[str]:
        """读取文件内容
        
        Args:
            file_path: 文件路径
            encoding: 文件编码，默认utf-8
            
        Returns:
            Optional[str]: 文件内容，如果读取失败返回None
        """
        try:
            path = Path(file_path)
            if not path.exists():
                self.logger.warning(f"文件不存在: {file_path}")
                return None
            
            with path.open('r', encoding=encoding) as f:
                content = f.read()
            
            self.logger.info(f"成功读取文件: {file_path}")
            return content
        
        except Exception as e:
            self.logger.error(f"读取文件 {file_path} 时发生错误: {str(e)}")
            return None
    
    def _build_markdown_prompt(self, file_config: Dict, search_results: Dict) -> str:
        """构建Markdown生成的提示词
        
        Args:
            file_config: 文件配置字典
            search_results: 搜索结果字典
            
        Returns:
            str: 生成提示词
        """
        try:
            # 提取搜索结果中的结构化数据
            structured_data = search_results.get('structured_results', [])
            
            # 构建提示词
            prompt = f"""
请根据以下信息生成一篇Markdown格式的文档：

[文档主题]
{file_config['topic']}

[文档用途]
{file_config['description']}

[内容结构]
{json.dumps(file_config.get('content_structure', {}), ensure_ascii=False, indent=2)}

[参考信息]
{json.dumps(structured_data, ensure_ascii=False, indent=2)}

要求：
1. 严格按照给定的内容结构组织文档
2. 使用Markdown语法格式化文档
3. 确保内容准确性，不要编造信息
4. 保持专业性和客观性
5. 适当使用图表、列表等形式增强可读性
"""
            return prompt
            
        except Exception as e:
            self.logger.error(f"构建Markdown提示词时发生错误: {str(e)}")
            return ""
    
    async def _generate_markdown_content(self, 
                                      file_config: Dict, 
                                      search_results: Dict,
                                      process_dir: Path) -> str:
        """生成Markdown格式的内容"""
        try:
            if not self.reasoning_engine:
                raise ValueError("未配置推理引擎，无法生成内容")
            
            # 保存生成提示词
            prompt = self._build_markdown_prompt(file_config, search_results)
            if not prompt:
                raise ValueError("生成提示词失败")
            
            prompt_path = process_dir / "markdown_prompt.md"
            prompt_path.write_text(prompt, encoding="utf-8")
            
            # 添加用户消息
            self.reasoning_engine.add_message("user", prompt)
            
            # 收集生成的内容
            full_response = ""
            process_path = process_dir / "generation_process.txt"
            
            async for chunk in self.reasoning_engine.get_stream_response(
                temperature=0.7,
                metadata={'stage': 'markdown_generation'}
            ):
                if chunk:
                    full_response += chunk
                    self.logger.info(f"\r生成Markdown内容: {full_response}")
                    # 保存生成过程
                    with process_path.open("a", encoding="utf-8") as f:
                        f.write(chunk)
            
            if not full_response:
                raise ValueError("生成内容为空")
            
            # 保存完整响应
            response_path = process_dir / "full_response.md"
            response_path.write_text(full_response, encoding="utf-8")
            
            # 提取Markdown内容
            # 1. 首先尝试提取第一个```markdown和最后一个```之间的内容
            start_marker = "```markdown\n"
            end_marker = "```"
            
            if start_marker in full_response:
                start_idx = full_response.find(start_marker) + len(start_marker)
                remaining_text = full_response[start_idx:]
                if end_marker in remaining_text:
                    end_idx = remaining_text.rfind(end_marker)
                    md_content = remaining_text[:end_idx].strip()
                    if md_content:
                        # 保存提取的Markdown内容
                        content_path = process_dir / "extracted_content.md"
                        content_path.write_text(md_content, encoding="utf-8")
                        return md_content
            
            # 2. 如果没有markdown标记，尝试提取任意代码块
            if "```" in full_response:
                start_idx = full_response.find("```") + 3
                # 跳过语言标识符所在行
                start_idx = full_response.find("\n", start_idx) + 1
                remaining_text = full_response[start_idx:]
                if "```" in remaining_text:
                    end_idx = remaining_text.rfind("```")
                    md_content = remaining_text[:end_idx].strip()
                    if md_content:
                        content_path = process_dir / "extracted_content.md"
                        content_path.write_text(md_content, encoding="utf-8")
                        return md_content
            
            # 3. 如果没有代码块，检查是否整个响应就是markdown
            if full_response.startswith('#') or full_response.startswith('---'):
                content_path = process_dir / "extracted_content.md"
                content_path.write_text(full_response, encoding="utf-8")
                return full_response
            
            # 如果都不符合，返回错误信息
            error_content = f"""# 生成失败

很抱歉，无法生成有效的Markdown内容。

原始响应：
{full_response[:200]}...
"""
            return error_content
            
        except Exception as e:
            self.logger.error(f"生成Markdown内容时发生错误: {str(e)}")
            error_info = {
                "timestamp": datetime.now().isoformat(),
                "error": str(e),
                "traceback": traceback.format_exc()
            }
            
            error_path = process_dir / "generation_error.json"
            with error_path.open("w", encoding="utf-8") as f:
                json.dump(error_info, f, ensure_ascii=False, indent=2)
            
            # 返回错误提示内容
            return f"""# 生成失败

很抱歉，生成内容时发生错误：

{str(e)}
"""
    
    async def _generate_html_content(self,
                                   file_config: Dict,
                                   search_results: Dict,
                                   process_dir: Path) -> str:
        """生成HTML格式的内容"""
        if not self.reasoning_engine:
            raise ValueError("未配置推理引擎，无法生成内容")
            
        # 从搜索结果中提取数据
        structured_data = search_results.get('structured', [])
        vector_data = search_results.get('vector_results', [])
        
        # 合并所有可用数据
        all_data = []
        
        # 处理结构化数据
        for item in structured_data:
            if isinstance(item.get('data'), str):
                try:
                    data_dict = json.loads(item['data'])
                    all_data.append({
                        'type': 'structured',
                        'file_name': item.get('_file_name', ''),
                        'content': data_dict.get('content', ''),
                        'timestamp': item.get('_processed_at', '')
                    })
                except json.JSONDecodeError:
                    continue
        
        # 处理向量数据
        for item in vector_data:
            if isinstance(item.get('data'), str):
                try:
                    data_dict = json.loads(item['data'])
                    all_data.append({
                        'type': 'vector',
                        'file_name': item.get('file_name', ''),
                        'content': data_dict.get('content', ''),
                        'similarity': item.get('similarity', 0)
                    })
                except json.JSONDecodeError:
                    continue
        
        if not all_data:
            raise ValueError("未找到可用的数据来生成HTML内容")
            
        # 生成HTML内容
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>{file_config.get('topic', 'AI分析报告')}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                .item {{ margin-bottom: 20px; padding: 10px; border: 1px solid #ddd; }}
                .structured {{ background-color: #f0f8ff; }}
                .vector {{ background-color: #fff0f5; }}
            </style>
        </head>
        <body>
            <h1>{file_config.get('topic', 'AI分析报告')}</h1>
            <div class="content">
        """
        
        for item in all_data:
            html_content += f"""
                <div class="item {item['type']}">
                    <h3>{item['file_name']}</h3>
                    <p>{item['content']}</p>
                    {'<p>相似度: {:.2f}</p>'.format(item['similarity']) if 'similarity' in item else ''}
                </div>
            """
        
        html_content += """
            </div>
        </body>
        </html>
        """
        
        # 保存生成的HTML
        html_path = process_dir / "content.html"
        html_path.write_text(html_content, encoding="utf-8")
        
        return html_content
    
    def _generate_error_html(self, error_message: str, title: str) -> str:
        """生成错误提示页面"""
        return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title} - 生成失败</title>
    <style>
        body {{
            font-family: system-ui, -apple-system, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            background: #f5f5f5;
        }}
        .error-container {{
            background: white;
            border-left: 5px solid #dc3545;
            padding: 20px;
            border-radius: 4px;
            box-shadow: 0 2px 4px rgba(0,0,0,0.1);
        }}
        h1 {{
            color: #dc3545;
            margin-top: 0;
        }}
    </style>
</head>
<body>
    <div class="error-container">
        <h1>内容生成失败</h1>
        <p>{error_message}</p>
        <p>请检查生成日志获取详细信息。</p>
    </div>
</body>
</html>"""
    
    async def _generate_csv_content(self,
                                  file_config: Dict,
                                  search_results: Dict,
                                  process_dir: Path) -> pd.DataFrame:
        """生成CSV格式的内容
        
        Args:
            file_config: 文件配置字典
            search_results: 搜索结果字典
            process_dir: 处理过程目录路径
            
        Returns:
            pd.DataFrame: 生成的CSV数据，如果生成失败返回空DataFrame
        """
        try:
            # 从搜索结果中提取结构化数据
            structured_data = search_results.get('structured', [])  # 改为 'structured'
            
            if not structured_data:
                self.logger.warning("未找到结构化数据，尝试从向量结果构建")
                vector_data = search_results.get('vector_results', [])
                if vector_data:
                    structured_data = vector_data
            
            if not structured_data:
                raise ValueError("未找到可用的结构化数据")
            
            # 将数据转换为DataFrame
            records = []
            for item in structured_data:
                if isinstance(item.get('data'), str):
                    try:
                        data_dict = json.loads(item['data'])
                        records.append({
                            'file_name': item.get('_file_name', ''),
                            'chunk_id': data_dict.get('chunk_id', ''),
                            'content': data_dict.get('content', '')
                        })
                    except json.JSONDecodeError:
                        continue
            
            df = pd.DataFrame(records)
            if not df.empty:
                # 保存原始CSV数据
                raw_path = process_dir / "raw_data.csv"
                df.to_csv(raw_path, index=False, encoding='utf-8-sig')
                return df
            
            raise ValueError("无法从结构化数据构建有效的DataFrame")
            
        except Exception as e:
            self.logger.error(f"生成CSV内容时发生错误: {str(e)}")
            # 记录详细错误信息
            error_info = {
                "timestamp": datetime.now().isoformat(),
                "error": str(e),
                "traceback": traceback.format_exc()
            }
            
            error_path = process_dir / "generation_error.json"
            with error_path.open("w", encoding="utf-8") as f:
                json.dump(error_info, f, ensure_ascii=False, indent=2)
            
            # 返回空DataFrame
            return pd.DataFrame()
    
    async def _generate_docx_content(self,
                                   file_config: Dict,
                                   search_results: Dict,
                                   process_dir: Path) -> bytes:
        """生成DOCX格式的内容（基于Markdown内容转换）"""
        try:
            # 先获取Markdown内容
            md_content = await self._generate_markdown_content(file_config, search_results, process_dir)
            
            # 创建Word文档
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            
            doc = Document()
            
            # 设置基本样式
            doc.styles['Normal'].font.name = '微软雅黑'
            doc.styles['Normal'].font.size = Pt(11)
            
            # 添加标题
            title = doc.add_heading(file_config['topic'], level=0)
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # 添加分隔线
            doc.add_paragraph().add_run('').add_break()
            
            # 解析Markdown内容
            current_list = []
            in_code_block = False
            
            for line in md_content.split('\n'):
                line = line.strip()
                
                # 跳过空行
                if not line:
                    doc.add_paragraph()
                    continue
                    
                # 处理代码块
                if line.startswith('```'):
                    in_code_block = not in_code_block
                    if in_code_block:
                        p = doc.add_paragraph()
                        p.style = 'Normal'
                        continue
                    else:
                        doc.add_paragraph()
                        continue
                
                if in_code_block:
                    p = doc.add_paragraph(line)
                    p.style = 'Normal'
                    continue
                
                # 处理标题
                if line.startswith('#'):
                    level = len(line.split()[0])  # 计算#的数量
                    text = line.lstrip('#').strip()
                    doc.add_heading(text, level=min(level, 9))
                    continue
                
                # 处理列表
                if line.startswith('- ') or line.startswith('* '):
                    text = line[2:].strip()
                    p = doc.add_paragraph(text)
                    p.style = 'List Bullet'
                    continue
                
                if line.startswith('1. ') or re.match(r'^\d+\. ', line):
                    text = line.split('. ', 1)[1].strip()
                    p = doc.add_paragraph(text)
                    p.style = 'List Number'
                    continue
                
                # 处理加粗文本
                if '**' in line:
                    p = doc.add_paragraph()
                    parts = line.split('**')
                    for i, part in enumerate(parts):
                        run = p.add_run(part)
                        if i % 2 == 1:  # 奇数索引表示加粗部分
                            run.bold = True
                    continue
                
                # 处理普通段落
                p = doc.add_paragraph(line)
                p.style = 'Normal'
            
            # 保存到字节流
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.read()
            
        except Exception as e:
            self.logger.error(f"生成DOCX内容时发生错误: {str(e)}")
            return b''
    
    async def generate_deliverables(self, 
                                  delivery_plan: Dict,
                                  search_results: Dict,
                                  delivery_config: Dict = None,
                                  test_mode: bool = False) -> List[str]:
        """生成交付文件
        
        Args:
            delivery_plan: 交付计划字典
            search_results: 搜索结果字典
            delivery_config: 交付配置字典
            test_mode: 是否为测试模式
            
        Returns:
            List[str]: 生成的文件路径列表
        """
        try:
            # 创建生成目录
            generation_dir = self.work_dir / "generations"
            generation_dir.mkdir(parents=True, exist_ok=True)
            
            # 保存输入数据
            input_dir = generation_dir / "inputs"
            input_dir.mkdir(exist_ok=True)
            
            with open(input_dir / "delivery_plan.json", "w", encoding="utf-8") as f:
                json.dump(delivery_plan, f, ensure_ascii=False, indent=2)
            
            with open(input_dir / "search_results.json", "w", encoding="utf-8") as f:
                json.dump(search_results, f, ensure_ascii=False, indent=2)
            
            if delivery_config:
                with open(input_dir / "delivery_config.json", "w", encoding="utf-8") as f:
                    json.dump(delivery_config, f, ensure_ascii=False, indent=2)
            
            # 创建输出目录
            output_dir = generation_dir / "outputs"
            output_dir.mkdir(exist_ok=True)
            
            # 创建生成过程记录目录
            process_dir = generation_dir / "process"
            process_dir.mkdir(exist_ok=True)
            
            generated_files = []
            
            # 获取交付文件配置
            delivery_files = delivery_config.get('delivery_files', {}) if delivery_config else {}
            
            if not delivery_files:  # 如果没有delivery_config，尝试从delivery_plan获取
                delivery_files = delivery_plan.get('delivery_files', {})
            
            # 生成每个交付文件
            for purpose, file_config in delivery_files.items():
                file_name = file_config['file_name']
                file_path = output_dir / file_name
                
                # 为每个文件创建单独的生成记录目录
                file_process_dir = process_dir / Path(file_name).stem
                file_process_dir.mkdir(exist_ok=True)
                
                try:
                    if test_mode:  # 测试模式生成占位文件
                        content = await self._generate_test_content(file_name)
                        self._save_test_content(content, file_path)
                        generated_files.append(str(file_path))
                        continue
                    
                    # 保存文件配置
                    with open(file_process_dir / "file_config.json", "w", encoding="utf-8") as f:
                        json.dump(file_config, f, ensure_ascii=False, indent=2)
                    
                    # 根据文件类型生成内容
                    content = None
                    if file_name.endswith('.md'):
                        content = await self._generate_markdown_content(
                            file_config, 
                            search_results,
                            file_process_dir
                        )
                        if content:
                            file_path.write_text(content, encoding='utf-8')
                            
                    elif file_name.endswith('.html'):
                        content = await self._generate_html_content(
                            file_config, 
                            search_results,
                            file_process_dir
                        )
                        if content:
                            file_path.write_text(content, encoding='utf-8')
                            
                    elif file_name.endswith('.csv'):
                        df = await self._generate_csv_content(
                            file_config, 
                            search_results,
                            file_process_dir
                        )
                        if not df.empty:
                            df.to_csv(file_path, index=False, encoding='utf-8-sig')
                            
                    elif file_name.endswith('.docx'):
                        content = await self._generate_docx_content(
                            file_config, 
                            search_results,
                            file_process_dir
                        )
                        if content:
                            file_path.write_bytes(content)
                    
                    if content is not None or (isinstance(content, pd.DataFrame) and not content.empty):
                        generated_files.append(str(file_path))
                        self.logger.info(f"已生成文件: {file_path}")
                        
                        # 记录生成成功
                        with open(file_process_dir / "generation_success.json", "w", encoding="utf-8") as f:
                            json.dump({
                                "timestamp": datetime.now().isoformat(),
                                "file_path": str(file_path),
                                "file_size": file_path.stat().st_size
                            }, f, ensure_ascii=False, indent=2)
                    else:
                        # 记录生成失败
                        with open(file_process_dir / "generation_failure.json", "w", encoding="utf-8") as f:
                            json.dump({
                                "timestamp": datetime.now().isoformat(),
                                "error": "生成内容为空"
                            }, f, ensure_ascii=False, indent=2)
                        
                except Exception as e:
                    self.logger.error(f"生成文件 {file_name} 时发生错误: {str(e)}")
                    # 记录错误信息
                    with open(file_process_dir / "generation_error.json", "w", encoding="utf-8") as f:
                        json.dump({
                            "timestamp": datetime.now().isoformat(),
                            "error": str(e),
                            "traceback": traceback.format_exc()
                        }, f, ensure_ascii=False, indent=2)
                    continue
            
            # 生成README文件
            readme_path = generation_dir / "README.md"
            self._generate_readme(
                readme_path,
                generated_files,
                delivery_files
            )
            
            # 保存生成统计信息
            stats = {
                "timestamp": datetime.now().isoformat(),
                "total_files": len(delivery_files),
                "generated_files": len(generated_files),
                "success_rate": len(generated_files) / len(delivery_files) if delivery_files else 0,
                "file_list": generated_files
            }
            
            with open(generation_dir / "generation_stats.json", "w", encoding="utf-8") as f:
                json.dump(stats, f, ensure_ascii=False, indent=2)
            
            return generated_files
            
        except Exception as e:
            self.logger.error(f"生成交付文件时发生错误: {str(e)}")
            raise 

    def _extract_file_paths(self, search_results: Dict) -> List[str]:
        """从搜索结果中提取所有不重复的文件路径
        
        Args:
            search_results: 搜索结果字典
            
        Returns:
            List[str]: 不重复的文件路径列表
        """
        try:
            file_paths = set()
            
            # 从结构化结果中提取
            structured_results = search_results.get('structured_results', [])
            for result in structured_results:
                if '_file_path' in result:
                    file_paths.add(result['_file_path'])
                    
            # 从向量结果中提取
            vector_results = search_results.get('vector_results', [])
            for result in vector_results:
                if 'file_path' in result:  # 注意向量结果中是file_path而不是_file_path
                    file_paths.add(result['file_path'])
            
            self.logger.debug(f"从搜索结果中提取了 {len(file_paths)} 个不重复文件路径")
            return list(file_paths)
            
        except Exception as e:
            self.logger.error(f"提取文件路径时发生错误: {str(e)}")
            return [] 

    def _save_qa_pairs(self, qa_pairs: Dict[str, str], qa_dir: Path) -> None:
        """保存问答对到文件
        
        Args:
            qa_pairs: 问答对字典
            qa_dir: QA对保存目录
        """
        try:
            # 保存完整的问答对
            qa_pairs_path = qa_dir / "all_qa_pairs.json"
            with qa_pairs_path.open('w', encoding='utf-8') as f:
                json.dump(qa_pairs, f, ensure_ascii=False, indent=2)
            
            # 为每个文件单独保存问答对
            for file_name, qa_content in qa_pairs.items():
                file_qa_path = qa_dir / f"{Path(file_name).stem}_qa.json"
                with file_qa_path.open('w', encoding='utf-8') as f:
                    json.dump({file_name: qa_content}, f, ensure_ascii=False, indent=2)
            
            self.logger.info(f"问答对已保存到目录: {qa_dir}")
            
        except Exception as e:
            self.logger.error(f"保存问答对时发生错误: {str(e)}")

    async def _create_qa_pairs(self, context: Dict) -> Optional[Dict[str, str]]:
        """读取所有源文件的内容，并生成问答对"""
        try:
            if not self.reasoning_engine:
                self.logger.warning("未配置推理引擎，无法生成内容总结")
                return None
            
            # 获取所有文件路径
            file_paths = self._extract_file_paths(context['search_results'])
            if not file_paths:
                self.logger.warning("未找到需要处理的源文件")
                return None
            
            delivery_plan = json.dumps(context['delivery_plan'], ensure_ascii=False, indent=2)
            
            # 存储所有文件的问答对
            all_qa_pairs = {}
            
            # 对每个文件单独处理
            for file_path in file_paths:
                content = self._read_file_content(file_path)
                if not content:
                    continue
                    
                file_name = Path(file_path).name
                
                # 构建针对单个文件的提示信息
                self.reasoning_engine.add_message("user", f"""
                    [file name]: {file_name}
                    [file content begin]
                    {content}
                    [file content end]
                    请从文件内容中提取出与delivery plan相关的内容，并生成问答对的形式。
                    [delivery plan begin]
                    {delivery_plan}
                    [delivery plan end]

                    问答对用JSON格式输出：
                    {{
                        "question": "问题",
                        "answer": "回答"
                    }}  
                """)
                
                # 使用流式输出收集完整响应
                file_qa_pairs = ""
                async for chunk in self.reasoning_engine.get_stream_response(
                    temperature=0.7,
                    metadata={'stage': 'single_file_summarization'}
                ):
                    file_qa_pairs += chunk
                    self.logger.info(f"\r生成{file_name}的问答对: {file_qa_pairs}")
                
                if file_qa_pairs:
                    all_qa_pairs[file_name] = file_qa_pairs
                    self.logger.info(f"成功生成{file_name}的问答对")
                else:
                    self.logger.warning(f"生成{file_name}的问答对失败")
            
            if not all_qa_pairs:
                self.logger.warning("未能成功生成任何文件的问答对")
                return None
            
            # 修改: 使用 work_dir 作为基础目录来保存 QA 对
            qa_dir = self.work_dir / "qa_pairs"
            qa_dir.mkdir(parents=True, exist_ok=True)
            self._save_qa_pairs(all_qa_pairs, qa_dir)
            
            return all_qa_pairs
            
        except Exception as e:
            self.logger.error(f"生成问答对时发生错误: {str(e)}")
            return None 

    def _generate_readme(self, readme_path: Path, generated_files: List[str], delivery_files: Dict) -> None:
        """生成README文件"""
        with readme_path.open('w', encoding='utf-8') as f:
            f.write("# 交付文件说明\n\n")
            f.write(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n\n")
            f.write("## 文件列表\n\n")
            for file_path in generated_files:
                file_name = Path(file_path).name
                # 查找文件用途，如果找不到则使用默认描述
                purpose = next((k for k, v in delivery_files.items() 
                             if v['file_name'] == file_name), None)
                f.write(f"- {file_name}\n")
                if purpose and purpose in delivery_files:
                    f.write(f"  - 用途：{purpose}\n")
                    f.write(f"  - 说明：{delivery_files[purpose]['description']}\n\n")
                else:
                    f.write("  - 用途：自动生成的补充文件\n")
                    f.write("  - 说明：根据主文件自动生成的配套文件\n\n")

    def _generate_test_content(self, file_name: str) -> str:
        """生成测试内容"""
        return f"# 测试文件\n这是测试模式生成的占位文件"

    def _save_test_content(self, content: str, file_path: Path) -> None:
        """保存测试内容到文件"""
        file_path.write_text(content, encoding='utf-8') 